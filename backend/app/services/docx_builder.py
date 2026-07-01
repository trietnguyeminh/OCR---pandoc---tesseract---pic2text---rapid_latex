"""DOCX assembly.

Two backends, chosen automatically:

* **pandoc** (preferred) — we serialise the document model to Markdown with
  ``$$...$$`` math and let pandoc emit *editable* Word equations (OMML). A
  generated reference.docx pins Times New Roman / Cambria Math and heading
  styles. Tables become real Word tables, figures are embedded.
* **python-docx** (fallback, always available) — builds the document directly;
  equations are rendered to images (matplotlib mathtext) so they are never lost.

Returns a stats dict the pipeline folds into the conversion report.
"""
from __future__ import annotations

import logging
import os
import re
import subprocess

from ..config import Settings, capabilities
from ..models import Block, BlockType, ConvertOptions, DocumentModel

logger = logging.getLogger("formudoc.docx")

_MD_SPECIAL = re.compile(r"([\\`*_{}\[\]#|<>])")
_EDITABLE_CONF = 0.45


def _esc(text: str) -> str:
    return _MD_SPECIAL.sub(r"\\\1", (text or "").replace("\n", " ")).strip()


def _img_url(path: str) -> str:
    # pandoc markdown: backslashes are escapes -> use forward slashes, and wrap
    # in <> so paths with spaces (e.g. "Minh Triet") are not split.
    return "<" + str(path).replace("\\", "/") + ">"


def _img_alt(text: str) -> str:
    return (text or "").replace("[", "(").replace("]", ")").replace("\n", " ")[:120]


# --------------------------------------------------------------------------- #
#  Reference document (fonts + styles) for the pandoc backend
# --------------------------------------------------------------------------- #
def build_reference_doc(path: str) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    for lvl, size in ((1, 18), (2, 15), (3, 13)):
        try:
            st = doc.styles[f"Heading {lvl}"]
            st.font.name = "Cambria"
            st.font.size = Pt(size)
        except KeyError:
            pass
    doc.save(path)
    return path


# --------------------------------------------------------------------------- #
#  Backend selection
# --------------------------------------------------------------------------- #
def _spans_to_openxml(md: str) -> str:
    """Convert AI color spans <span style="color:#RRGGBB">text</span> into inline
    raw-openxml runs so pandoc emits real coloured runs in the .docx. Text that
    contains math/code/newlines is left as plain text (color dropped) to avoid
    breaking LaTeX conversion."""
    import re, html as _html
    pat = re.compile(
        r'<span\s+style\s*=\s*["\'][^"\']*color\s*:\s*#?([0-9A-Fa-f]{6})[^"\']*["\']\s*>(.*?)</span>',
        re.S | re.I)

    def repl(m):
        hexv, text = m.group(1), m.group(2)
        if ("$" in text) or ("`" in text) or ("\n" in text) or ("<" in text):
            return text  # keep plain; don't break math/code
        t = _html.escape(text, quote=False)
        return (f'`<w:r><w:rPr><w:color w:val="{hexv.upper()}"/></w:rPr>'
                f'<w:t xml:space="preserve">{t}</w:t></w:r>`{{=openxml}}')
    return pat.sub(repl, md or "")


def markdown_to_docx(markdown: str, out_path: str, settings: Settings) -> None:
    """Render a Markdown string (text + $...$ math + tables) to .docx via pandoc.
    Used by the Pix2Text full-page path."""
    import pypandoc
    import shutil
    import tempfile
    work = tempfile.mkdtemp(prefix="formudoc_md_")
    try:
        ref = build_reference_doc(os.path.join(work, "reference.docx"))
        pypandoc.convert_text(_spans_to_openxml(markdown) or "(empty)", "docx", format="markdown",
                              outputfile=out_path,
                              extra_args=["--reference-doc", ref, "--wrap=preserve"])
    finally:
        shutil.rmtree(work, ignore_errors=True)


def _write_md_sidecar(document, options, out_path, stats) -> None:
    try:
        md = build_markdown(document, options)
        md_path = os.path.splitext(out_path)[0] + ".md"
        with open(md_path, "w", encoding="utf-8") as fh:
            fh.write(md)
        stats["md_path"] = md_path
        logger.info("Wrote Markdown sidecar %s", md_path)
    except Exception as exc:  # pragma: no cover
        logger.warning("Could not write Markdown sidecar: %s", exc)


def build_markdown(document: DocumentModel, options: ConvertOptions) -> str:
    """A clean text+LaTeX Markdown of the document (no images). Saved as a .md
    sidecar so the user can read/edit/re-process the text separately."""
    parts: list[str] = []
    last_page = None
    for b in document.blocks:
        if b.type in (BlockType.header, BlockType.footer):
            continue
        if options.preserve_layout and last_page is not None and b.page != last_page:
            parts.append("\n---\n")          # page separator
        last_page = b.page
        t = (b.text or "").replace("\n", " ").strip()
        if b.type == BlockType.title:
            parts.append(f"# {t}")
        elif b.type == BlockType.heading:
            parts.append(f"{'#' * min(max(b.level, 1) + 1, 6)} {t}")
        elif b.type == BlockType.caption:
            parts.append(f"*{t}*")
        elif b.type == BlockType.formula:
            if b.latex:
                parts.append(f"$$\n{b.latex}\n$$")
            elif t:
                parts.append(f"`{t}`")
        elif b.type == BlockType.figure:
            parts.append("*(hình / figure)*")
        elif b.type == BlockType.table:
            if b.table_data:
                parts.append(_table_to_md(b))
        else:
            if t:
                parts.append(t)
    return "\n\n".join(p for p in parts if p and p.strip())


def build_docx(document: DocumentModel, options: ConvertOptions, out_path: str,
               settings: Settings) -> dict:
    caps = capabilities()
    if caps.get("pandoc"):
        try:
            stats = _build_with_pandoc(document, options, out_path, settings)
            _write_md_sidecar(document, options, out_path, stats)
            return stats
        except Exception as exc:  # pragma: no cover
            logger.warning("Pandoc backend failed (%s); falling back to python-docx", exc)
    stats = _build_with_python_docx(document, options, out_path, settings)
    _write_md_sidecar(document, options, out_path, stats)
    return stats


# --------------------------------------------------------------------------- #
#  Pandoc backend
# --------------------------------------------------------------------------- #
def _table_to_md(block: Block) -> str:
    grid = block.table_data or []
    if not grid:
        return ""
    n_cols = max(len(r) for r in grid)
    rows = [[(c or "").replace("\n", " ").replace("|", "\\|") for c in r] +
            [""] * (n_cols - len(r)) for r in grid]
    header = rows[0]
    body = rows[1:] or [[""] * n_cols]
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join(["---"] * n_cols) + " |"]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _pagebreak() -> str:
    return '```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```'


def _build_with_pandoc(document, options, out_path, settings) -> dict:
    import pypandoc
    import shutil
    import tempfile

    stats = {"backend": "pandoc", "editable_equations": 0,
             "image_fallback_equations": 0, "tables_emitted": 0,
             "figures_emitted": 0, "removed_headers_footers": 0}

    # Stage every referenced image into a clean work dir and reference it by a
    # simple ascii basename. This avoids ALL pandoc path pitfalls on Windows
    # (backslashes are markdown escapes; spaces split image URLs).
    work = tempfile.mkdtemp(prefix="formudoc_md_")
    _imgmap: dict[str, str] = {}

    def _stage(src: str | None) -> str | None:
        if not src or not os.path.exists(src):
            return None
        if src in _imgmap:
            return _imgmap[src]
        ext = os.path.splitext(src)[1] or ".png"
        name = f"img{len(_imgmap)}{ext}"
        try:
            shutil.copyfile(src, os.path.join(work, name))
        except Exception:
            return None
        _imgmap[src] = name
        return name

    parts: list[str] = []
    last_page = None
    for b in document.blocks:
        if b.type in (BlockType.header, BlockType.footer):
            stats["removed_headers_footers"] += 1
            continue
        if options.preserve_layout and last_page is not None and b.page != last_page:
            parts.append(_pagebreak())
        last_page = b.page

        if b.type == BlockType.title:
            parts.append(f"# {_esc(b.text)}")
        elif b.type == BlockType.heading:
            parts.append(f"{'#' * min(max(b.level, 1) + 1, 6)} {_esc(b.text)}")
        elif b.type == BlockType.caption:
            parts.append(f"*{_esc(b.text)}*")
        elif b.type == BlockType.footnote:
            parts.append(f"<sub>{_esc(b.text)}</sub>")
        elif b.type == BlockType.formula:
            name = _stage(b.image_path)
            if b.editable_equation and b.latex:
                # Mathpix returns markdown with $..$ math -> emit as a paragraph;
                # pix2tex/heuristic return pure LaTeX -> wrap as a display equation.
                parts.append(b.latex if "$" in b.latex else f"$$ {b.latex} $$")
                stats["editable_equations"] += 1
            elif name:
                parts.append(f"![]({name})")
                stats["image_fallback_equations"] += 1
            elif b.latex:
                parts.append(b.latex if "$" in b.latex else f"$$ {b.latex} $$")
                stats["editable_equations"] += 1
        elif b.type == BlockType.figure:
            name = _stage(b.image_path)
            if name:
                parts.append(f"![]({name})")
                stats["figures_emitted"] += 1
        elif b.type == BlockType.table:
            if (b.confidence or 0) >= 0.5 and b.table_data:
                parts.append(_table_to_md(b))
                stats["tables_emitted"] += 1
            else:
                name = _stage(b.image_path)
                if name:
                    parts.append(f"![]({name})")
                    stats["tables_emitted"] += 1
                elif b.table_data:
                    parts.append(_table_to_md(b))
                    stats["tables_emitted"] += 1

        else:  # paragraph / list_item
            parts.append(_esc(b.text))

    markdown = "\n\n".join(p for p in parts if p)
    ref = build_reference_doc(os.path.join(work, "reference.docx"))
    extra = ["--reference-doc", ref, "--wrap=preserve", "--resource-path", work]
    try:
        pypandoc.convert_text(markdown, "docx", format="markdown",
                              outputfile=out_path, extra_args=extra)
    finally:
        shutil.rmtree(work, ignore_errors=True)
    logger.info("Pandoc produced %s (%d editable eq, %d image eq, %d figures)",
                out_path, stats["editable_equations"],
                stats["image_fallback_equations"], stats["figures_emitted"])
    return stats


# --------------------------------------------------------------------------- #
#  python-docx fallback backend
# --------------------------------------------------------------------------- #
def _render_latex_png(latex: str, path: str) -> bool:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig = plt.figure(figsize=(0.01, 0.01))
        fig.text(0, 0, f"${latex}$", fontsize=16)
        fig.savefig(path, dpi=200, bbox_inches="tight", pad_inches=0.05,
                    transparent=True)
        plt.close(fig)
        return os.path.exists(path)
    except Exception as exc:
        logger.debug("mathtext render failed for %r: %s", latex, exc)
        return False


def _build_with_python_docx(document, options, out_path, settings) -> dict:
    from docx import Document
    from docx.shared import Inches, Pt

    stats = {"backend": "python-docx", "editable_equations": 0,
             "image_fallback_equations": 0, "tables_emitted": 0,
             "figures_emitted": 0, "removed_headers_footers": 0}

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    last_page = None

    for b in document.blocks:
        if b.type in (BlockType.header, BlockType.footer):
            stats["removed_headers_footers"] += 1
            continue
        if options.preserve_layout and last_page is not None and b.page != last_page:
            doc.add_page_break()
        last_page = b.page

        if b.type == BlockType.title:
            doc.add_heading(b.text, level=0)
        elif b.type == BlockType.heading:
            doc.add_heading(b.text, level=min(max(b.level, 1), 4))
        elif b.type == BlockType.caption:
            p = doc.add_paragraph(); r = p.add_run(b.text); r.italic = True
        elif b.type == BlockType.footnote:
            p = doc.add_paragraph(); r = p.add_run(b.text); r.font.size = Pt(9)
        elif b.type == BlockType.formula:
            png = str(settings.asset_dir / f"eq_{abs(hash((b.page, b.text)))}.png")
            if b.latex and _render_latex_png(b.latex, png):
                doc.add_picture(png, width=Inches(min(5.5, 0.12 * len(b.latex) + 1)))
                stats["image_fallback_equations"] += 1
            elif b.image_path and os.path.exists(b.image_path):
                doc.add_picture(b.image_path, width=Inches(4))
                stats["image_fallback_equations"] += 1
            else:
                p = doc.add_paragraph(); r = p.add_run(b.latex or b.text)
                r.font.name = "Cambria Math"
                stats["image_fallback_equations"] += 1
        elif b.type == BlockType.figure:
            if b.image_path and os.path.exists(b.image_path):
                try:
                    doc.add_picture(b.image_path, width=Inches(5))
                    stats["figures_emitted"] += 1
                except Exception:
                    pass
        elif b.type == BlockType.table:
            if b.table_data and (b.confidence or 0) >= 0.5:
                _add_docx_table(doc, b.table_data)
                stats["tables_emitted"] += 1
            elif b.image_path and os.path.exists(b.image_path):
                doc.add_picture(b.image_path, width=Inches(5.5))
                stats["tables_emitted"] += 1
            elif b.table_data:
                _add_docx_table(doc, b.table_data)
                stats["tables_emitted"] += 1
        else:
            for line in (b.text or "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

    doc.save(out_path)
    logger.info("python-docx produced %s", out_path)
    return stats


def _add_docx_table(doc, grid):
    n_cols = max(len(r) for r in grid)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    for row in grid:
        cells = table.add_row().cells
        for i in range(n_cols):
            cells[i].text = (row[i] if i < len(row) else "") or ""
